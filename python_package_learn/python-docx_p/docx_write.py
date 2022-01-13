from docx import Document

document = Document()
# 添加标题默认为1
document.add_heading('The role of dolphins', level=2)
# 添加分页，强制换页
document.add_page_break()
"""
表中数据修改
"""
# 添加表
table = document.add_table(rows=2, cols=2)
# 获取表的单元格
cell = table.cell(0, 1)
cell.text = 'parrot, possibly dead'  # 修改单元格数据
print(cell.text)  # 获取单元格的数据
# 获取表中的所有行
# 为所有行组成的列表,可使用for row in table.rows
rows = table.rows
row = table.rows[1]
row.cells[0].text = 'Foo bar to you.'
row.cells[1].text = 'And a hearty foo bar to you too sir!'
table.add_row()  # 添加行
# 获取表中的列
columns = table.columns
# 设置表格的样式
table.style = 'LightShading-Accent1'
"""
添加图片
"""
from docx.shared import Inches

document.add_picture('image-filename.png', width=Inches(1.0))
"""
添加段落
"""
# 添加段落
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
# 添加样式
# 样式名称与在Word用户界面（UI）中显示的名称完全相同
paragraph.style = 'List Bullet'
paragraph = document.add_paragraph('Normal text, ')

# 在当前段落继续添加文字
run = paragraph.add_run('dolor')
# 添加文字为加粗或斜体
# bold为加粗
# 斜体
# 方式１
paragraph.add_run('dolor2').bold = True
# 方式2
# is equivalent to:
run2 = paragraph.add_run('dolor2')
run2.bold = True
# 为继续添加文字添加样式,下面的这两种方式相同
# 方式１
paragraph.add_run('text with emphasis.', 'Emphasis')
# 方式2
paragraph = document.add_paragraph('Normal text, ')
run3 = paragraph.add_run('text with emphasis.')
run3.style = 'Emphasis'

# 文档保存

document.save('test.docx')
